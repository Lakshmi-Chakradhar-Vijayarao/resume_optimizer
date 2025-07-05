import streamlit as st
import time
from io import BytesIO
from docx import Document

from openai_helper import call_openai
from resume_to_text import extract_text_from_file
from prompts import PROMPT_TEMPLATES  # Load all prompts

# --- Streamlit page setup ---
st.set_page_config(page_title="Resume Optimizer", layout="centered")
st.title("📄 AI Resume Optimizer")
st.markdown("Step-by-step: Enter Role ➡️ Paste JD ➡️ Upload Resume ➡️ Download Optimized Version")

# --- Inputs ---
role = st.text_input("🎯 Job Role", placeholder="e.g., Full Stack Developer")
jd_text = st.text_area("📋 Paste Job Description")
uploaded_file = st.file_uploader("📎 Upload Resume (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])

# --- Prompt steps ---
steps = [
    "spot_flaws",
    "rewrite_for_impact",
    "ats_boost",
    "craft_hook",
    "upgrade_experience",
    "format_fix",
    "tailor_for_role",
    "benchmark_me"
]

# --- DOCX generation function ---
def generate_docx_resume(content: str) -> BytesIO:
    doc = Document()
    for line in content.strip().split('\n'):
        if line.strip() == "---":
            doc.add_paragraph("")  # line break
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Main logic ---
if role and jd_text and uploaded_file:
    resume_text = extract_text_from_file(uploaded_file)

    if resume_text:
        with st.spinner("⏳ Optimizing your resume..."):
            final_outputs = {}
            for step in steps:
                prompt_template = PROMPT_TEMPLATES[step]
                # Handle prompts with and without JD
                if "{jd}" in prompt_template:
                    prompt = prompt_template.format(role=role, jd=jd_text, resume=resume_text)
                else:
                    prompt = prompt_template.format(role=role, resume=resume_text)
                try:
                    output = call_openai(prompt)
                    final_outputs[step] = output
                except Exception as e:
                    final_outputs[step] = f"⚠️ Error: {str(e)}"
                time.sleep(25)  # Rate limit buffer

        # Generate readable full text
        optimized_resume = "\n\n".join(
            [f"🔹 {step.replace('_', ' ').title()}:\n{final_outputs[step]}" for step in steps]
        )

        # Extract the best final version for docx (from rewrite_for_impact)
        final_resume_text = final_outputs.get("rewrite_for_impact", optimized_resume)
        docx_file = generate_docx_resume(final_resume_text)

        # --- Output UI ---
        st.subheader("✅ Optimized Resume")
        st.text_area("Preview:", value=final_resume_text, height=400)

        st.download_button(
            label="📥 Download Final Resume (.docx)",
            data=docx_file,
            file_name="Optimized_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        if st.button("🔁 Start New Optimization"):
            st.experimental_rerun()

    else:
        st.error("❌ Couldn't extract text from the uploaded file.")
else:
    st.info("📌 Please fill in all 3 fields to start: Role, JD, and Resume.")
